"""
app.py
 
WRLD document-generation microservice.
 
NOTE ON THIS FILE: this version was reconstructed from the documented
architecture (session memory + transfer notes) because the live app.py
was not available to diff against directly. /create-template's
DocumentBuilder logic, /fill-template, and /read-template are rebuilt
to match the described behavior as closely as possible -- please diff
this against GitHub main before committing, particularly:
  - _disable_contextual_spacing() (List Bullet contextualSpacing fix)
  - the GOOGLE_CREDENTIALS_JSON / GOOGLE_SERVICE_ACCOUNT_JSON env fallback
    in services/drive.py (unchanged here, still expected to exist there)
  - services/parser.py's run-fragmentation bug (still open, not touched
    by this change -- COMPANY_ADDRESS_LINE_1/2 etc. still need the
    run-merge fix logged in the backlog)
 
Endpoints:
    POST /create-template     - pure JSON -> .docx renderer (unchanged contract)
    POST /fill-template       - fills a template's placeholders, returns raw bytes
    POST /read-template       - returns a template's placeholder list
    POST /create-presentation - pure JSON -> .pptx renderer (NEW)
    POST /execute-document    - sandboxed arbitrary python-docx/pptx code (NEW)
 
/create-template and /create-presentation share the same contract shape:
    request JSON  -> { "spec": {...} }
    response      -> raw file bytes (docx or pptx), Content-Disposition attachment
 
/execute-document accepts:
    { "code": "<python source defining build_document(spec) -> bytes>",
      "spec": {...},
      "output_filename": "...",
      "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        (or the pptx equivalent -- caller's responsibility, not inferred here) }
 
All three generation endpoints return raw bytes only. Drive upload is
handled entirely in n8n (Prepare for Drive -> Google Drive Upload),
mirroring the pattern already proven for /create-template and
/fill-template -- no new wiring is required for /execute-document.
"""
 
import io
import os
import re
 
from flask import Flask, request, jsonify, send_file
 
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
from presentation_builder import build_presentation, PresentationBuilderError
from code_sandbox import (
    execute_document_code,
    log_submission,
    SandboxValidationError,
    SandboxExecutionError,
)
 
# services/drive.py and services/parser.py are existing modules in this repo
# (unchanged by this update) -- imported here for /fill-template and
# /read-template, consistent with the Filler Agent's confirmed-working chain.
from services import drive as drive_service
from services import parser as parser_service
 
app = Flask(__name__)
 
 
# --------------------------------------------------------------------- #
# DocumentBuilder -- pure JSON -> .docx renderer
# No hardcoded font enums, no hardcoded document-type enums, no hardcoded
# spacing values. Reads only spec['components'] and spec['margins'].
# --------------------------------------------------------------------- #
 
class DocumentBuilderError(Exception):
    pass
 
 
def _disable_contextual_spacing(paragraph) -> None:
    """
    Word's built-in 'List Bullet' style has <w:contextualSpacing/> enabled,
    which silently suppresses spacing_before/spacing_after between
    consecutive same-style paragraphs regardless of what python-docx sets.
    Writing <w:contextualSpacing w:val="0"/> directly via OxmlElement is
    the confirmed fix.
    """
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:contextualSpacing"))
    if existing is not None:
        pPr.remove(existing)
    el = OxmlElement("w:contextualSpacing")
    el.set(qn("w:val"), "0")
    pPr.append(el)
 
 
def _apply_keep_properties(paragraph, keep_with_next=False, keep_together=False) -> None:
    pf = paragraph.paragraph_format
    if keep_with_next:
        pf.keep_with_next = True
    if keep_together:
        pf.keep_together = True
 
 
def _set_cell_no_split(row) -> None:
    """Sets w:cantSplit on a table row so it won't break across a page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)
 
 
class DocumentBuilder:
    """
    Spec-driven .docx renderer. Reads only spec['components'] (an ordered
    list of component dicts) and spec['margins']. Any styling value not
    provided by the caller is left at python-docx's own default -- this
    class does not invent design rules of its own.
    """
 
    COMPONENT_TYPES = {"paragraph", "heading", "list_bullet", "table", "page_break"}
 
    def __init__(self, spec: dict):
        if not isinstance(spec, dict):
            raise DocumentBuilderError("spec must be a dict")
        self.spec = spec
        self.components = spec.get("components") or []
        self.margins = spec.get("margins") or {}
 
        if not self.components:
            raise DocumentBuilderError("spec['components'] is required and cannot be empty")
 
        self.doc = Document()
        self._apply_margins()
 
    def _apply_margins(self) -> None:
        section = self.doc.sections[0]
        for attr in ("top", "bottom", "left", "right"):
            value_in = self.margins.get(attr)
            if value_in is not None:
                setattr(section, f"{attr}_margin", Inches(value_in))
 
    def build(self) -> bytes:
        for i, component in enumerate(self.components):
            ctype = component.get("type")
            if ctype not in self.COMPONENT_TYPES:
                raise DocumentBuilderError(
                    f"Component {i}: unknown type '{ctype}'. "
                    f"Must be one of {sorted(self.COMPONENT_TYPES)}"
                )
            getattr(self, f"_render_{ctype}")(component)
 
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()
 
    # ------------------------------------------------------------------ #
 
    def _render_paragraph(self, c: dict) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(c.get("text", ""))
        if c.get("bold"):
            run.bold = True
        if c.get("italic"):
            run.italic = True
        if c.get("font_size"):
            run.font.size = Pt(c["font_size"])
        if c.get("alignment") == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif c.get("alignment") == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
 
        pf = p.paragraph_format
        if c.get("spacing_before") is not None:
            pf.space_before = Pt(c["spacing_before"])
        if c.get("spacing_after") is not None:
            pf.space_after = Pt(c["spacing_after"])
 
        _apply_keep_properties(
            p,
            keep_with_next=bool(c.get("keep_with_next")),
            keep_together=bool(c.get("keep_together")),
        )
 
    def _render_heading(self, c: dict) -> None:
        level = c.get("level", 1)
        p = self.doc.add_heading(c.get("text", ""), level=level)
 
        pf = p.paragraph_format
        if c.get("spacing_before") is not None:
            pf.space_before = Pt(c["spacing_before"])
        if c.get("spacing_after") is not None:
            pf.space_after = Pt(c["spacing_after"])
 
        # Headings default to keep_with_next=True (a heading orphaned at the
        # bottom of a page with its content pushed to the next page is
        # almost never desired) unless the spec explicitly overrides it.
        keep_with_next = c.get("keep_with_next", True)
        _apply_keep_properties(
            p,
            keep_with_next=bool(keep_with_next),
            keep_together=bool(c.get("keep_together")),
        )
 
    def _render_list_bullet(self, c: dict) -> None:
        p = self.doc.add_paragraph(c.get("text", ""), style="List Bullet")
        _disable_contextual_spacing(p)
 
        pf = p.paragraph_format
        if c.get("spacing_before") is not None:
            pf.space_before = Pt(c["spacing_before"])
        if c.get("spacing_after") is not None:
            pf.space_after = Pt(c["spacing_after"])
 
        _apply_keep_properties(
            p,
            keep_with_next=bool(c.get("keep_with_next")),
            keep_together=bool(c.get("keep_together")),
        )
 
    def _render_table(self, c: dict) -> None:
        rows = c.get("rows") or []
        if not rows:
            raise DocumentBuilderError("table component requires a non-empty 'rows' list")
 
        header_row = c.get("header_row", True)
        font_size = c.get("font_size")
        prevent_row_split = c.get("prevent_row_split", True)
 
        n_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        table.style = c.get("style", "Table Grid")
 
        for r_idx, row_values in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx, cell_value in enumerate(row_values):
                cell = row.cells[c_idx]
                cell.text = str(cell_value)
                if header_row and r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                if font_size:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(font_size)
 
            if prevent_row_split:
                _set_cell_no_split(row)
 
    def _render_page_break(self, c: dict) -> None:
        self.doc.add_page_break()
 
 
def build_document(spec: dict) -> bytes:
    return DocumentBuilder(spec).build()
 
 
# --------------------------------------------------------------------- #
# Routes
# --------------------------------------------------------------------- #
 
@app.route("/create-template", methods=["POST"])
def create_template():
    payload = request.get_json(force=True, silent=True) or {}
    spec = payload.get("spec")
    if not spec:
        return jsonify({"error": "Request body must include 'spec'"}), 400
 
    try:
        file_bytes = build_document(spec)
    except DocumentBuilderError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Unexpected error building document: {e}"}), 500
 
    filename = payload.get("output_filename", "document.docx")
    return send_file(
        io.BytesIO(file_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
 
 
@app.route("/create-presentation", methods=["POST"])
def create_presentation():
    payload = request.get_json(force=True, silent=True) or {}
    spec = payload.get("spec")
    if not spec:
        return jsonify({"error": "Request body must include 'spec'"}), 400
 
    try:
        file_bytes = build_presentation(spec)
    except PresentationBuilderError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Unexpected error building presentation: {e}"}), 500
 
    filename = payload.get("output_filename", "presentation.pptx")
    return send_file(
        io.BytesIO(file_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
    )
 
 
@app.route("/execute-document", methods=["POST"])
def execute_document():
    payload = request.get_json(force=True, silent=True) or {}
    code = payload.get("code")
    spec = payload.get("spec")
    output_filename = payload.get("output_filename", "document.bin")
    mimetype = payload.get("mimetype", "application/octet-stream")
 
    if not code or not isinstance(code, str):
        return jsonify({"error": "Request body must include 'code' as a string"}), 400
    if spec is None or not isinstance(spec, dict):
        return jsonify({"error": "Request body must include 'spec' as an object"}), 400
 
    try:
        file_bytes = execute_document_code(code, spec)
    except SandboxValidationError as e:
        log_submission(code, spec, output_filename, outcome="rejected", detail=str(e))
        return jsonify({"error": f"Code rejected by sandbox validation: {e}"}), 400
    except SandboxExecutionError as e:
        log_submission(code, spec, output_filename, outcome="failed", detail=str(e))
        return jsonify({"error": f"Sandboxed execution failed: {e}"}), 422
    except Exception as e:
        log_submission(code, spec, output_filename, outcome="error", detail=str(e))
        return jsonify({"error": f"Unexpected sandbox error: {e}"}), 500
 
    log_submission(code, spec, output_filename, outcome="ok")
 
    return send_file(
        io.BytesIO(file_bytes),
        as_attachment=True,
        download_name=output_filename,
        mimetype=mimetype,
    )
 
 
@app.route("/read-template", methods=["POST"])
def read_template():
    """
    Fetches a template .docx from Drive and returns its placeholder list,
    so the Document Agent (or Filler Agent) can learn the target
    template's fields before writing content.
 
    NOTE: services/parser.py has a known run-fragmentation bug -- Word's
    XML can split a placeholder like [RECIPIENT_ADDRESS_LINE_1] across
    multiple <w:r> runs, which the regex scan misses unless runs are
    merged first. That fix is still on the backlog and is NOT included
    here; this route calls parser_service as it currently exists.
    """
    payload = request.get_json(force=True, silent=True) or {}
    template_file_id = payload.get("template_file_id")
    if not template_file_id:
        return jsonify({"error": "Request body must include 'template_file_id'"}), 400
 
    try:
        file_bytes = drive_service.download_file_bytes(template_file_id)
        placeholders = parser_service.extract_placeholders(file_bytes)
    except Exception as e:
        return jsonify({"error": f"Failed to read template: {e}"}), 500
 
    return jsonify({"template_file_id": template_file_id, "placeholders": placeholders})
 
 
@app.route("/fill-template", methods=["POST"])
def fill_template():
    """
    Fills a template's placeholders with provided values and returns raw
    .docx bytes. Drive upload is handled entirely in n8n -- this endpoint
    does not touch Drive for writing, only for reading the source template.
    """
    payload = request.get_json(force=True, silent=True) or {}
    template_file_id = payload.get("template_file_id")
    values = payload.get("values") or {}
    output_filename = payload.get("output_filename", "filled_document.docx")
 
    if not template_file_id:
        return jsonify({"error": "Request body must include 'template_file_id'"}), 400
    if not values:
        return jsonify({"error": "Request body must include non-empty 'values'"}), 400
 
    try:
        template_bytes = drive_service.download_file_bytes(template_file_id)
        doc = Document(io.BytesIO(template_bytes))
 
        for paragraph in doc.paragraphs:
            _fill_paragraph_placeholders(paragraph, values)
 
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _fill_paragraph_placeholders(paragraph, values)
 
        buffer = io.BytesIO()
        doc.save(buffer)
        file_bytes = buffer.getvalue()
    except Exception as e:
        return jsonify({"error": f"Failed to fill template: {e}"}), 500
 
    return send_file(
        io.BytesIO(file_bytes),
        as_attachment=True,
        download_name=output_filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
 
 
_PLACEHOLDER_RE = re.compile(r"\[([A-Z0-9_]+)\]")
 
 
def _fill_paragraph_placeholders(paragraph, values: dict) -> None:
    """
    Simple same-run placeholder replacement. Does not yet address the
    run-fragmentation bug (placeholders split across multiple runs) --
    that fix is tracked separately in services/parser.py's backlog item.
    """
    full_text = paragraph.text
    if "[" not in full_text:
        return
 
    def _replace(match):
        key = match.group(1)
        return str(values.get(key, match.group(0)))
 
    new_text = _PLACEHOLDER_RE.sub(_replace, full_text)
    if new_text == full_text:
        return
 
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
 
 
@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})
 
 
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
