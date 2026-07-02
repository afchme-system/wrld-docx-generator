import re
import logging
from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)

class TemplateParser:
    """Parse .docx templates to extract placeholders and structure"""
    
    PLACEHOLDER_PATTERN = r'\[([A-Z_]+)\]'
    
    @staticmethod
    def extract_placeholders(docx_bytes):
        """Extract all bracket placeholders from a .docx file"""
        try:
            doc = Document(BytesIO(docx_bytes))
            placeholders = set()
            
            # Scan all paragraphs
            for para in doc.paragraphs:
                for run in para.runs:
                    matches = re.findall(TemplateParser.PLACEHOLDER_PATTERN, run.text)
                    for match in matches:
                        placeholders.add(f"[{match}]")
            
            # Scan all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                matches = re.findall(TemplateParser.PLACEHOLDER_PATTERN, run.text)
                                for match in matches:
                                    placeholders.add(f"[{match}]")
            
            placeholders_list = sorted(list(placeholders))
            logger.info(f"Found {len(placeholders_list)} placeholders: {placeholders_list}")
            return placeholders_list
        except Exception as e:
            logger.error(f"Failed to extract placeholders: {e}")
            raise
    
    @staticmethod
    def analyze_structure(docx_bytes):
        """Analyze document structure (sections, tables, etc.)"""
        try:
            doc = Document(BytesIO(docx_bytes))
            
            structure = {
                "page_count": len(doc.sections),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "has_letterhead": False,
                "sections": []
            }
            
            # Detect common sections by content
            section_names = []
            for i, para in enumerate(doc.paragraphs[:10]):  # Check first 10 paragraphs
                text = para.text.strip().lower()
                if 'date' in text or '[DATE]' in para.text:
                    section_names.append('header')
                elif 'recipient' in text or '[RECIPIENT' in para.text:
                    section_names.append('recipient_block')
                elif 'dear' in text or 'hello' in text:
                    section_names.append('salutation')
            
            if section_names:
                structure['sections'] = list(set(section_names))
            
            # Check for letterhead (images in first section)
            first_section = doc.sections[0] if doc.sections else None
            if first_section:
                structure['has_letterhead'] = len(doc.paragraphs) > 0 and doc.paragraphs[0].style.name in ['Header']
            
            logger.info(f"Document structure: {structure}")
            return structure
        except Exception as e:
            logger.error(f"Failed to analyze structure: {e}")
            raise
    
    @staticmethod
    def get_metadata(docx_bytes, file_name=''):
        """Get comprehensive template metadata"""
        try:
            doc = Document(BytesIO(docx_bytes))
            
            # Core properties
            core_props = doc.core_properties
            
            metadata = {
                "file_name": file_name,
                "title": core_props.title or '',
                "subject": core_props.subject or '',
                "author": core_props.author or '',
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "placeholders": TemplateParser.extract_placeholders(docx_bytes),
                "structure": TemplateParser.analyze_structure(docx_bytes)
            }
            
            logger.info(f"Template metadata: {metadata['title']} with {len(metadata['placeholders'])} placeholders")
            return metadata
        except Exception as e:
            logger.error(f"Failed to get metadata: {e}")
            raise
