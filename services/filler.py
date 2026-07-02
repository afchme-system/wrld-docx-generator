import re
import logging
from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)

class TemplateFiller:
    """Fill .docx templates with values, preserving formatting via run-level replacement"""
    
    PLACEHOLDER_PATTERN = r'\[([A-Z_]+)\]'
    
    @staticmethod
    def fill_template(docx_bytes, replacements):
        """
        Fill a .docx template with provided replacements.
        
        Args:
            docx_bytes: Binary content of .docx file
            replacements: Dict mapping [PLACEHOLDER] -> value
        
        Returns:
            Tuple of (filled_doc_bytes, report)
            report contains: replacements_made, unmatched_placeholders, extra_keys
        """
        try:
            doc = Document(BytesIO(docx_bytes))
            
            # Track operations
            replacements_made = 0
            found_placeholders = set()
            report = {
                "replacements_made": 0,
                "unmatched_placeholders": [],
                "extra_keys": [],
                "warnings": []
            }
            
            # Extract all placeholders in document
            for para in doc.paragraphs:
                TemplateFiller._fill_paragraph_runs(para, replacements, found_placeholders)
                replacements_made += len([p for p in found_placeholders])
            
            # Fill tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            TemplateFiller._fill_paragraph_runs(para, replacements, found_placeholders)
            
            # Check for unmatched placeholders in document
            all_doc_placeholders = TemplateFiller._extract_all_placeholders(doc)
            unmatched = [p for p in all_doc_placeholders if p not in found_placeholders]
            
            # Check for extra keys (not in document)
            extra_keys = [k for k in replacements.keys() if k not in all_doc_placeholders]
            
            report['replacements_made'] = len(found_placeholders)
            report['unmatched_placeholders'] = unmatched
            report['extra_keys'] = extra_keys
            
            if unmatched:
                msg = f"Unmatched placeholders: {unmatched}"
                logger.warning(msg)
                report['warnings'].append(msg)
            
            if extra_keys:
                msg = f"Extra keys not in template: {extra_keys}"
                logger.warning(msg)
                report['warnings'].append(msg)
            
            # Save to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info(f"Filled template: {len(found_placeholders)} replacements made")
            return output.getvalue(), report
        
        except Exception as e:
            logger.error(f"Failed to fill template: {e}", exc_info=True)
            raise
    
    @staticmethod
    def _fill_paragraph_runs(para, replacements, found_placeholders):
        """
        Fill placeholders in a paragraph by replacing at RUN level.
        
        Critical: Replacing at run level preserves formatting (bold, italic, font, size, color)
        Replacing at paragraph.text level destroys all formatting.
        """
        for run in para.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    # Direct match: entire run text is the placeholder
                    if run.text == placeholder:
                        run.text = str(value)
                        found_placeholders.add(placeholder)
                    # Partial match: placeholder is part of run text
                    elif placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                        found_placeholders.add(placeholder)
    
    @staticmethod
    def _extract_all_placeholders(doc):
        """Extract all placeholders currently in document"""
        placeholders = set()
        
        # Paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                matches = re.findall(TemplateFiller.PLACEHOLDER_PATTERN, run.text)
                for match in matches:
                    placeholders.add(f"[{match}]")
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            matches = re.findall(TemplateFiller.PLACEHOLDER_PATTERN, run.text)
                            for match in matches:
                                placeholders.add(f"[{match}]")
        
        return placeholders
