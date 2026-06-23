import json
import io
import logging
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from google.oauth2.service_account import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

logger = logging.getLogger(__name__)

class TemplateFiller:
    """Handles downloading templates from Google Drive and filling placeholders"""
    
    def __init__(self, drive_folder_id, google_credentials_json):
        """
        Initialize with Google Drive credentials
        
        Args:
            drive_folder_id: Google Drive folder ID containing templates
            google_credentials_json: JSON string with service account credentials
        """
        self.drive_folder_id = drive_folder_id
        self.drive_service = self._init_drive_service(google_credentials_json)
        self.template_cache = {}  # Cache template IDs to avoid repeated searches
    
    def _init_drive_service(self, credentials_json_str):
        """Initialize Google Drive API service"""
        try:
            if not credentials_json_str:
                raise ValueError(
                    "GOOGLE_CREDENTIALS_JSON environment variable not set. "
                    "Please set this to your Google service account JSON credentials."
                )
            
            try:
                credentials_dict = json.loads(credentials_json_str)
            except json.JSONDecodeError as e:
                raise ValueError(
                    f"Invalid JSON in GOOGLE_CREDENTIALS_JSON: {str(e)}. "
                    "Ensure you're passing the entire JSON as a string."
                )
            
            credentials = Credentials.from_service_account_info(
                credentials_dict,
                scopes=['https://www.googleapis.com/auth/drive.readonly']
            )
            service = build('drive', 'v3', credentials=credentials)
            logger.info("Google Drive API service initialized successfully")
            return service
        except Exception as e:
            logger.error(f"Failed to initialize Drive service: {str(e)}", exc_info=True)
            raise
    
    def _find_template(self, template_name):
        """
        Find template file ID by name in Drive folder
        
        Args:
            template_name: Name of template file (e.g., "WRLD-Template-Letter (1).docx")
        
        Returns:
            File ID string
        
        Raises:
            FileNotFoundError if template not found
        """
        # Check cache first
        if template_name in self.template_cache:
            return self.template_cache[template_name]
        
        try:
            query = (
                f"name = '{template_name}' and "
                f"'{self.drive_folder_id}' in parents and "
                f"trashed = false"
            )
            results = self.drive_service.files().list(
                q=query,
                spaces='drive',
                fields='files(id, name)',
                pageSize=1
            ).execute()
            
            files = results.get('files', [])
            if not files:
                raise FileNotFoundError(
                    f"Template '{template_name}' not found in Drive folder {self.drive_folder_id}"
                )
            
            file_id = files[0]['id']
            self.template_cache[template_name] = file_id
            logger.info(f"Found template: {template_name} ({file_id})")
            return file_id
        
        except Exception as e:
            logger.error(f"Error finding template: {str(e)}")
            raise
    
    def _download_template(self, file_id):
        """
        Download template file from Drive
        
        Args:
            file_id: Google Drive file ID
        
        Returns:
            Binary content of .docx file
        """
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            logger.info(f"Downloaded template: {file_id}")
            return file_content.getvalue()
        
        except Exception as e:
            logger.error(f"Error downloading template: {str(e)}")
            raise
    
    def _replace_placeholders(self, doc, content):
        """
        Replace [Placeholders] in document with values from content dict
        
        Args:
            doc: python-docx Document object
            content: Dictionary with key-value pairs
        
        Returns:
            Modified Document object
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in content.items():
                value_str = str(value) if value is not None else ""
                
                # Match [Placeholder] format
                placeholder = f"[{key}]"
                
                if placeholder in paragraph.text:
                    # Replace in paragraph text
                    self._replace_text_in_paragraph(paragraph, placeholder, value_str)
                    logger.debug(f"Replaced {placeholder} in paragraph")
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in content.items():
                            value_str = str(value) if value is not None else ""
                            placeholder = f"[{key}]"
                            
                            if placeholder in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, placeholder, value_str)
                                logger.debug(f"Replaced {placeholder} in table cell")
        
        return doc
    
    def _replace_text_in_paragraph(self, paragraph, placeholder, replacement):
        """
        Replace placeholder in a paragraph while preserving formatting
        
        Args:
            paragraph: python-docx Paragraph object
            placeholder: String to find (e.g., "[Date]")
            replacement: String to replace with
        """
        if placeholder not in paragraph.text:
            return
        
        # If placeholder is in a single run, simple replacement
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = paragraph.runs[0].text.replace(placeholder, replacement)
            return
        
        # If placeholder spans multiple runs, rebuild the paragraph
        full_text = paragraph.text
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, replacement)
            
            # Clear existing runs
            for run in paragraph.runs:
                r = run._element
                r.getparent().remove(r)
            
            # Add new run with replaced text
            paragraph.add_run(new_text)
    
    def fill_template(self, template_name, content):
        """
        Fill a template with content and return binary .docx
        
        Args:
            template_name: Name of template file
            content: Dictionary with placeholder:value pairs
                    e.g., {"date": "22 June 2026", "recipient_full_name": "John Smith"}
        
        Returns:
            Binary content of filled .docx file
        """
        try:
            logger.info(f"Starting fill_template: {template_name}")
            
            # Find and download template
            file_id = self._find_template(template_name)
            template_binary = self._download_template(file_id)
            
            # Load into python-docx
            doc = Document(io.BytesIO(template_binary))
            logger.debug(f"Loaded template into Document object")
            
            # Replace placeholders
            doc = self._replace_placeholders(doc, content)
            logger.info(f"Replaced {len(content)} placeholders")
            
            # Save to binary
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info(f"Successfully filled template: {template_name}")
            return output.getvalue()
        
        except Exception as e:
            logger.error(f"Error filling template: {str(e)}")
            raise
    
    def list_available_templates(self):
        """
        List all available templates in the Drive folder
        
        Returns:
            List of template names
        """
        try:
            query = (
                f"'{self.drive_folder_id}' in parents and "
                f"mimeType = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' and "
                f"trashed = false"
            )
            results = self.drive_service.files().list(
                q=query,
                spaces='drive',
                fields='files(id, name)',
                pageSize=10
            ).execute()
            
            templates = [f['name'] for f in results.get('files', [])]
            logger.info(f"Listed {len(templates)} templates")
            return templates
        
        except Exception as e:
            logger.error(f"Error listing templates: {str(e)}")
            raise
