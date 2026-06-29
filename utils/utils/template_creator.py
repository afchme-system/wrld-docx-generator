import io
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class TemplateCreator:
    """Create new Word documents from JSON specifications"""
    
    def __init__(self):
        pass
    
    def create_document(self, title, sections, placeholders, metadata_fields=None, 
                       cover_page=True, header=True, footer=True, page_numbers=True):
        """
        Create a new Word document from specification.
        
        Args:
            title: Document title
            sections: List of section dicts with 'name' and optional 'placeholder'
            placeholders: List of placeholder strings (e.g., "[INVESTOR_NAME]")
            metadata_fields: Dict of document metadata
            cover_page: Whether to add a cover page
            header: Whether to add header
            footer: Whether to add footer
            page_numbers: Whether to add page numbers
        
        Returns:
            Binary .docx file content
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add cover page if requested
        if cover_page:
            self._add_cover_page(doc, title, metadata_fields or {})
            doc.add_page_break()
        
        # Add table of contents placeholder
        doc.add_heading('Contents', level=1)
        doc.add_paragraph('[Table of Contents will be generated here]')
        doc.add_page_break()
        
        # Add sections
        for section in sections:
            section_name = section.get('name', 'Untitled Section')
            section_placeholder = section.get('placeholder', '')
            
            doc.add_heading(section_name, level=1)
            
            if section_placeholder:
                doc.add_paragraph(section_placeholder, style='Normal')
            else:
                doc.add_paragraph('[Content for ' + section_name + ']', style='Normal')
            
            doc.add_paragraph()  # spacing
        
        # Add placeholders section
        if placeholders:
            doc.add_page_break()
            doc.add_heading('Placeholders', level=1)
            for placeholder in placeholders:
                doc.add_paragraph(f'• {placeholder}', style='List Bullet')
        
        # Add header if requested
        if header:
            self._add_header(doc, title)
        
        # Add footer if requested
        if footer:
            self._add_footer(doc, page_numbers)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"Created document: {title}, {len(sections)} sections, {len(placeholders)} placeholders")
        return doc_bytes.getvalue()
    
    def _add_cover_page(self, doc, title, metadata):
        """Add a professional cover page"""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Metadata
        if metadata:
            for key, value in metadata.items():
                meta_para = doc.add_paragraph(f'{key}: {value}')
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if meta_para.runs:
                    meta_para.runs[0].font.size = Pt(11)
        
        # Date placeholder
        doc.add_paragraph()
        date_para = doc.add_paragraph('[Document Date: _______________]')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if date_para.runs:
            date_para.runs[0].font.italic = True
    
    def _add_header(self, doc, title):
        """Add header to all sections"""
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = title
        if header_para.runs:
            header_para.runs[0].font.size = Pt(10)
            header_para.runs[0].font.italic = True
    
    def _add_footer(self, doc, page_numbers=True):
        """Add footer to all sections"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        
        if page_numbers:
            footer_para.text = "Page [#] of [##]"
        else:
            footer_para.text = "---"
        
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_para.runs:
            footer_para.runs[0].font.size = Pt(10)
